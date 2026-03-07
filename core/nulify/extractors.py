"""
Text Extraction Engine — Extract text content from multiple file formats.
Supports: PDF, DOCX, TXT, CSV, SQL, JSON, Images (PNG, JPG, JPEG, BMP, TIFF, WEBP)
"""

import os
import csv
import json
import logging

logger = logging.getLogger(__name__)

# Image file types supported by OCR
IMAGE_TYPES = {'png', 'jpg', 'jpeg', 'bmp', 'tiff', 'webp'}


def extract_text(file_path, file_type):
    """
    Main dispatcher: extract text from a file based on its type.
    Returns the full text content as a string.
    """
    file_type = file_type.lower().strip('.')

    extractors = {
        'pdf': extract_from_pdf,
        'docx': extract_from_docx,
        'txt': extract_from_txt,
        'csv': extract_from_csv,
        'xlsx': extract_from_xlsx,
        'sql': extract_from_sql,
        'json': extract_from_json,
    }

    # Image types → OCR extraction (text only, no boxes)
    if file_type in IMAGE_TYPES:
        text, _boxes = extract_from_image(file_path)
        return text

    extractor = extractors.get(file_type)
    if not extractor:
        raise ValueError(f"Unsupported file type: {file_type}")

    return extractor(file_path)


def extract_from_pdf(file_path):
    """Extract text and properly formatted tables from PDF using pdfplumber."""
    try:
        import pdfplumber
    except ImportError:
        return _extract_from_pdf_fallback(file_path)

    text_parts = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            tables = page.find_tables()
            table_extractor_settings = {}
            if not tables:
                # Try text alignment heuristic for borderless tables
                table_extractor_settings = {"vertical_strategy": "text", "horizontal_strategy": "text"}
                tables = page.find_tables(table_settings=table_extractor_settings)
            
            if not tables:
                text = page.extract_text(layout=True)
                if text:
                    import re
                    processed_lines = []
                    for line in text.split('\n'):
                        stripped_line = line.strip()
                        # If a line has gaps of 3 or more spaces between words, it's a borderless column
                        if re.search(r'   +', stripped_line):
                            table_line = re.sub(r'   +', ' | ', stripped_line)
                            table_line = table_line.strip('|').strip()
                            processed_lines.append(f"| {table_line} |")
                        else:
                            processed_lines.append(stripped_line)
                    text_parts.append('\n'.join(processed_lines))
                continue
                
            # Filter out text that belongs to a table
            table_bboxes = [t.bbox for t in tables]
            
            def not_in_table(obj):
                if 'x0' not in obj or 'top' not in obj:
                    return True
                for (x0, top, x1, bottom) in table_bboxes:
                    # check if objection center is inside table bbox
                    cx = (obj['x0'] + obj['x1']) / 2
                    cy = (obj['top'] + obj['bottom']) / 2
                    if x0 <= cx <= x1 and top <= cy <= bottom:
                        return False
                return True

            non_table_page = page.filter(not_in_table)
            text = non_table_page.extract_text()
            if text:
                text_parts.append(text)
                
            for table in page.extract_tables(table_settings=table_extractor_settings):
                for row in table:
                    clean_row = [str(cell).strip().replace('\n', ' ') if cell else "" for cell in row]
                    if any(clean_row):
                        text_parts.append(' | '.join(clean_row))
                text_parts.append('')

    return '\n'.join(text_parts)


def _extract_from_pdf_fallback(file_path):
    """Fallback text extraction if pdfplumber is not installed."""
    try:
        from PyPDF2 import PdfReader
    except ImportError:
        raise ImportError("pdfplumber or PyPDF2 is required for PDF processing.")

    reader = PdfReader(file_path)
    text_parts = []
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text_parts.append(page_text)
    return '\n'.join(text_parts)


def extract_from_docx(file_path):
    """Extract text from DOCX using python-docx."""
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx is required for DOCX processing. Install with: pip install python-docx")

    doc = Document(file_path)
    text_parts = []

    # Extract from paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)

    # Extract from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                text_parts.append(' | '.join(row_text))

    return '\n'.join(text_parts)


def extract_from_txt(file_path):
    """Extract text from plain text files."""
    encodings = ['utf-8', 'latin-1', 'cp1252', 'ascii']
    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                return f.read()
        except (UnicodeDecodeError, UnicodeError):
            continue
    raise ValueError("Unable to decode the text file with supported encodings")


def extract_from_csv(file_path):
    """Extract text from CSV files."""
    text_parts = []
    encodings = ['utf-8', 'latin-1', 'cp1252']

    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding, newline='') as f:
                reader = csv.reader(f)
                for row in reader:
                    text_parts.append(' | '.join(str(cell) for cell in row))
            return '\n'.join(text_parts)
        except (UnicodeDecodeError, UnicodeError):
            continue

    raise ValueError("Unable to decode the CSV file with supported encodings")


def extract_from_xlsx(file_path):
    """Extract text from Excel (.xlsx) files using openpyxl."""
    try:
        from openpyxl import load_workbook
    except ImportError:
        raise ImportError("openpyxl is required for XLSX processing. Install with: pip install openpyxl")

    wb = load_workbook(file_path, read_only=True, data_only=True)
    text_parts = []

    for sheet in wb.worksheets:
        if len(wb.worksheets) > 1:
            text_parts.append(f"--- Sheet: {sheet.title} ---")
        for row in sheet.iter_rows(values_only=True):
            cells = [str(cell) if cell is not None else '' for cell in row]
            if any(c.strip() for c in cells):
                text_parts.append(' | '.join(cells))

    wb.close()
    return '\n'.join(text_parts)


def extract_from_sql(file_path):
    """Extract text from SQL dump files."""
    return extract_from_txt(file_path)


def extract_from_json(file_path):
    """Extract text from JSON files by flattening all values."""
    with open(file_path, 'r', encoding='utf-8') as f:
        data = json.load(f)

    text_parts = []
    _flatten_json(data, text_parts)
    return '\n'.join(text_parts)


def _flatten_json(obj, parts, prefix=''):
    """Recursively flatten JSON and collect all string values."""
    if isinstance(obj, dict):
        for key, value in obj.items():
            _flatten_json(value, parts, f"{prefix}{key}: ")
    elif isinstance(obj, list):
        for item in obj:
            _flatten_json(item, parts, prefix)
    else:
        parts.append(f"{prefix}{obj}")


# ══════════════════════════════════════════════════════════════════════
#  IMAGE / OCR EXTRACTION
# ══════════════════════════════════════════════════════════════════════

def _configure_tesseract():
    """Auto-detect Tesseract binary on Windows."""
    import shutil
    if shutil.which('tesseract'):
        return  # Already on PATH

    # Common Windows install paths
    common_paths = [
        r'C:\Program Files\Tesseract-OCR\tesseract.exe',
        r'C:\Program Files (x86)\Tesseract-OCR\tesseract.exe',
        os.path.expanduser(r'~\AppData\Local\Tesseract-OCR\tesseract.exe'),
    ]
    for path in common_paths:
        if os.path.isfile(path):
            import pytesseract
            pytesseract.pytesseract.tesseract_cmd = path
            logger.info(f"Tesseract found at: {path}")
            return

    logger.warning("Tesseract OCR not found. Image processing may fail.")


def extract_from_image(file_path):
    """
    Extract text and word-level bounding boxes from an image using OCR.
    Returns (text, boxes) where boxes is a list of dicts:
        [{'text': str, 'x': int, 'y': int, 'w': int, 'h': int, 'conf': float}]
    """
    try:
        import pytesseract
        from PIL import Image
    except ImportError:
        raise ImportError(
            "pytesseract and Pillow are required for image processing. "
            "Install with: pip install pytesseract Pillow"
        )

    _configure_tesseract()

    img = Image.open(file_path)

    # Get word-level data with bounding boxes
    data = pytesseract.image_to_data(img, output_type=pytesseract.Output.DICT)

    boxes = []
    text_parts = []
    current_line = -1
    line_words = []

    for i in range(len(data['text'])):
        word = data['text'][i].strip()
        conf = int(data['conf'][i]) if data['conf'][i] != '-1' else 0
        line_num = data['line_num'][i]

        if word and conf > 20:  # Filter low-confidence noise
            boxes.append({
                'text': word,
                'x': data['left'][i],
                'y': data['top'][i],
                'w': data['width'][i],
                'h': data['height'][i],
                'conf': conf / 100.0,
            })

            if line_num != current_line:
                if line_words:
                    text_parts.append(' '.join(line_words))
                line_words = [word]
                current_line = line_num
            else:
                line_words.append(word)

    # Flush last line
    if line_words:
        text_parts.append(' '.join(line_words))

    full_text = '\n'.join(text_parts)
    return full_text, boxes


def extract_text_with_boxes(file_path, file_type):
    """
    Wrapper that returns (text, boxes) for images, or (text, None) for other files.
    Used by the upload view to get bounding box data for image sanitization.
    """
    file_type = file_type.lower().strip('.')

    if file_type in IMAGE_TYPES:
        return extract_from_image(file_path)

    # Non-image: standard extraction, no boxes
    text = extract_text(file_path, file_type)
    return text, None
