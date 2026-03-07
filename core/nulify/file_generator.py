"""
File Generator — Generate sanitized output files and PDF reports.
Supports in-place PDF redaction using PyMuPDF to preserve original formatting.
"""

import logging
from io import BytesIO
from django.core.files.base import ContentFile

logger = logging.getLogger(__name__)


def generate_sanitized_file(uploaded_file, sanitized_text, detections=None, method='redaction'):
    """
    Create a sanitized file in the same format as the original.
    For PDFs: uses PyMuPDF to redact PII in-place, preserving tables and layout.
    Returns (filename, ContentFile) tuple for saving to SanitizedFile.sanitized_file.
    """
    base_name = uploaded_file.original_filename.rsplit('.', 1)[0]
    file_type = uploaded_file.file_type.lower()

    # ── DOCX output ──
    if file_type == 'docx':
        try:
            from docx import Document
            doc = Document(uploaded_file.file.path)
            # In-place replacement in paragraphs
            if detections:
                from .sanitizer import _get_replacement
                for para in doc.paragraphs:
                    for d in detections:
                        original_val = d['value']
                        if original_val in para.text:
                            replacement = _get_replacement(original_val, d['type'], method)
                            for run in para.runs:
                                if original_val in run.text:
                                    run.text = run.text.replace(original_val, replacement)
                # Also handle tables in DOCX
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for d in detections:
                                original_val = d['value']
                                if original_val in cell.text:
                                    replacement = _get_replacement(original_val, d['type'], method)
                                    for para in cell.paragraphs:
                                        for run in para.runs:
                                            if original_val in run.text:
                                                run.text = run.text.replace(original_val, replacement)
            else:
                # Fallback: rebuild from sanitized text
                doc = Document()
                for line in sanitized_text.split('\n'):
                    doc.add_paragraph(line)
            buf = BytesIO()
            doc.save(buf)
            buf.seek(0)
            filename = f"sanitized_{uploaded_file.id}_{base_name}.docx"
            return filename, ContentFile(buf.read())
        except ImportError:
            pass  # Fall through to text output

    # ── PDF output (in-place redaction with PyMuPDF) ──
    if file_type == 'pdf':
        result = _generate_pdf_inplace(uploaded_file, detections, method, base_name)
        if result:
            return result
        # Fallback to ReportLab rebuild if PyMuPDF fails
        result = _generate_pdf_reportlab(uploaded_file, sanitized_text, base_name)
        if result:
            return result

    # ── SQL output ──
    if file_type == 'sql':
        filename = f"sanitized_{uploaded_file.id}_{base_name}.sql"
        content = ContentFile(sanitized_text.encode('utf-8'))
        return filename, content

    # ── XLSX output ──
    if file_type == 'xlsx':
        try:
            from openpyxl import Workbook
            wb = Workbook()
            ws = wb.active
            ws.title = "Sanitized"
            for line in sanitized_text.split('\n'):
                if line.startswith('--- Sheet:') and line.endswith('---'):
                    continue  # Skip sheet header markers
                if '|' in line:
                    cells = [cell.strip() for cell in line.split('|')]
                else:
                    cells = [line]
                ws.append(cells)
            buf = BytesIO()
            wb.save(buf)
            buf.seek(0)
            filename = f"sanitized_{uploaded_file.id}_{base_name}.xlsx"
            return filename, ContentFile(buf.read())
        except ImportError:
            pass  # Fall through to text output

    # ── CSV output ──
    if file_type == 'csv':
        filename = f"sanitized_{uploaded_file.id}_{base_name}.csv"
        content = ContentFile(sanitized_text.encode('utf-8'))
        return filename, content

    # ── JSON output ──
    if file_type == 'json':
        filename = f"sanitized_{uploaded_file.id}_{base_name}.json"
        content = ContentFile(sanitized_text.encode('utf-8'))
        return filename, content

    # ── Default: TXT output ──
    filename = f"sanitized_{uploaded_file.id}_{base_name}.txt"
    content = ContentFile(sanitized_text.encode('utf-8'))
    return filename, content

def _generate_pdf_inplace(uploaded_file, detections, method, base_name):
    """
    Use PyMuPDF (fitz) to redact PII directly on the original PDF.
    This preserves ALL original formatting: tables, borders, images, fonts.
    """
    if not detections:
        # No PII found — just copy the original
        try:
            with open(uploaded_file.file.path, 'rb') as f:
                filename = f"sanitized_{uploaded_file.id}_{base_name}.pdf"
                return filename, ContentFile(f.read())
        except Exception:
            return None

    try:
        import fitz  # PyMuPDF
        from .sanitizer import _get_replacement
    except ImportError:
        logger.warning("PyMuPDF not installed, falling back to ReportLab")
        return None

    try:
        doc = fitz.open(uploaded_file.file.path)

        # Build replacement map: original_value → replacement_value
        replacements = {}
        for d in detections:
            original_val = d['value']
            if original_val not in replacements:
                replacements[original_val] = _get_replacement(original_val, d['type'], method)

        # Sort by length descending to replace longer matches first
        sorted_originals = sorted(replacements.keys(), key=len, reverse=True)

        for page in doc:
            for original_val in sorted_originals:
                replacement_val = replacements[original_val]
                # Search for all instances of this PII value on this page
                text_instances = page.search_for(original_val)
                for inst in text_instances:
                    # Add redaction annotation at the exact position
                    page.add_redact_annot(inst, text=replacement_val, fontsize=0,
                                         fill=(1, 1, 1))  # white background
                # Apply all redactions on this page
                if text_instances:
                    page.apply_redactions()

        buf = BytesIO()
        doc.save(buf)
        doc.close()
        buf.seek(0)
        filename = f"sanitized_{uploaded_file.id}_{base_name}.pdf"
        return filename, ContentFile(buf.read())

    except Exception as e:
        logger.error(f"PyMuPDF in-place redaction failed: {e}")
        return None


def _generate_pdf_reportlab(uploaded_file, sanitized_text, base_name):
    """Fallback: rebuild PDF from sanitized text using ReportLab."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import cm
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet

        buf = BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=A4,
                                topMargin=2*cm, bottomMargin=2*cm,
                                leftMargin=2*cm, rightMargin=2*cm)
        styles = getSampleStyleSheet()
        story = []
        for line in sanitized_text.split('\n'):
            escaped = line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            story.append(Paragraph(escaped, styles['Normal']))
            story.append(Spacer(1, 4))
        doc.build(story)
        buf.seek(0)
        filename = f"sanitized_{uploaded_file.id}_{base_name}.pdf"
        return filename, ContentFile(buf.read())
    except ImportError:
        return None


def generate_report_pdf(uploaded_file, detections, sanitized_file=None):
    """
    Generate a PDF sanitization report.
    Returns a BytesIO buffer containing the PDF.
    """
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import inch, cm
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, HRFlowable
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
    except ImportError:
        # Fallback: generate a simple text report
        return _generate_text_report(uploaded_file, detections, sanitized_file)

    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4,
                            topMargin=1.5 * cm, bottomMargin=1.5 * cm,
                            leftMargin=2 * cm, rightMargin=2 * cm)

    styles = getSampleStyleSheet()

    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle', parent=styles['Title'],
        fontSize=22, textColor=colors.HexColor('#0B132B'),
        spaceAfter=6
    )
    subtitle_style = ParagraphStyle(
        'CustomSubtitle', parent=styles['Normal'],
        fontSize=11, textColor=colors.HexColor('#64748B'),
        spaceAfter=20
    )
    heading_style = ParagraphStyle(
        'CustomHeading', parent=styles['Heading2'],
        fontSize=14, textColor=colors.HexColor('#7B61FF'),
        spaceBefore=20, spaceAfter=10
    )
    info_style = ParagraphStyle(
        'InfoStyle', parent=styles['Normal'],
        fontSize=10, leading=16
    )

    story = []

    # ── Header ──
    story.append(Paragraph("Nullify — Sanitization Report", title_style))
    story.append(Paragraph("PII Detection & Data Sanitization Platform", subtitle_style))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#E2E8F0')))
    story.append(Spacer(1, 15))

    # ── File Information ──
    story.append(Paragraph("File Information", heading_style))

    file_info = [
        ['Property', 'Details'],
        ['File Name', uploaded_file.original_filename],
        ['File Type', uploaded_file.file_type.upper()],
        ['File Size', _format_size(uploaded_file.file_size)],
        ['Upload Date', uploaded_file.uploaded_at.strftime('%d %B %Y, %I:%M %p')],
        ['Uploaded By', uploaded_file.uploaded_by.username],
        ['Risk Score', f"{uploaded_file.risk_score}% ({uploaded_file.risk_level().title()})"],
        ['Sanitization Method', sanitized_file.get_method_display() if sanitized_file else 'N/A'],
        ['Total PII Found', str(len(detections))],
    ]

    t = Table(file_info, colWidths=[150, 300])
    t.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#7B61FF')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
        ('TOPPADDING', (0, 0), (-1, 0), 10),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#E2E8F0')),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F8FAFC')]),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('LEFTPADDING', (0, 0), (-1, -1), 10),
    ]))
    story.append(t)
    story.append(Spacer(1, 20))

    # ── PII Detection Summary ──
    story.append(Paragraph("PII Detection Summary", heading_style))

    pii_summary = {}
    for d in detections:
        pii_type = d.pii_type if hasattr(d, 'pii_type') else d.get('type', 'Unknown')
        pii_summary[pii_type] = pii_summary.get(pii_type, 0) + 1

    if pii_summary:
        summary_data = [['PII Type', 'Count', 'Risk Weight']]
        from .pii_engine import PII_WEIGHTS
        for pii_type, count in sorted(pii_summary.items()):
            weight = PII_WEIGHTS.get(pii_type, 1)
            summary_data.append([pii_type, str(count), str(weight)])

        t2 = Table(summary_data, colWidths=[180, 100, 120])
        t2.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#00C2A8')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 10),
            ('TOPPADDING', (0, 0), (-1, 0), 10),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#E2E8F0')),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#F8FAFC')]),
            ('ALIGN', (1, 0), (-1, -1), 'CENTER'),
            ('LEFTPADDING', (0, 0), (-1, -1), 10),
        ]))
        story.append(t2)
    else:
        story.append(Paragraph("No PII detected in this file.", info_style))

    story.append(Spacer(1, 30))

    # ── Footer ──
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#E2E8F0')))
    story.append(Spacer(1, 8))
    footer_style = ParagraphStyle(
        'Footer', parent=styles['Normal'],
        fontSize=8, textColor=colors.HexColor('#94A3B8'),
        alignment=TA_CENTER
    )
    story.append(Paragraph(
        "Generated by Nullify — PII Detection & Data Sanitization Platform",
        footer_style
    ))

    doc.build(story)
    buffer.seek(0)
    return buffer


def _generate_text_report(uploaded_file, detections, sanitized_file):
    """Fallback text-based report if reportlab is not available."""
    buffer = BytesIO()
    lines = [
        "=" * 60,
        "NULLIFY — SANITIZATION REPORT",
        "=" * 60,
        "",
        f"File Name:      {uploaded_file.original_filename}",
        f"File Type:      {uploaded_file.file_type.upper()}",
        f"Upload Date:    {uploaded_file.uploaded_at}",
        f"Risk Score:     {uploaded_file.risk_score}%",
        f"Method:         {sanitized_file.get_method_display() if sanitized_file else 'N/A'}",
        f"Total PII:      {len(detections)}",
        "",
        "-" * 60,
        "PII DETECTION SUMMARY",
        "-" * 60,
    ]

    pii_summary = {}
    for d in detections:
        pii_type = d.pii_type if hasattr(d, 'pii_type') else d.get('type', 'Unknown')
        pii_summary[pii_type] = pii_summary.get(pii_type, 0) + 1

    for pii_type, count in sorted(pii_summary.items()):
        lines.append(f"  {pii_type:20s} : {count}")

    lines.extend(["", "=" * 60])
    buffer.write('\n'.join(lines).encode('utf-8'))
    buffer.seek(0)
    return buffer


def _format_size(size_bytes):
    """Format file size in human-readable format."""
    if size_bytes < 1024:
        return f"{size_bytes} B"
    elif size_bytes < 1024 * 1024:
        return f"{size_bytes / 1024:.1f} KB"
    else:
        return f"{size_bytes / (1024 * 1024):.1f} MB"
